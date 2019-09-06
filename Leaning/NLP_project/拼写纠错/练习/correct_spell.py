from docx import Document
from nltk import sent_tokenize,word_tokenize
from docx.shared import RGBColor
from correct_relu import correct_text_generic

from string import punctuation

COUNT_CORRECT = 0

file = Document('../data/Spelling Error.docx')

print('段落数量：',len(file.paragraphs))

# 创建新的文件
document = Document()

def write_correct_paragraph(i):
    global COUNT_CORRECT

    paragraph = file.paragraphs[i].text.strip()

    sentences = sent_tokenize(text=paragraph)

    word_lists = [word_tokenize(sentence) for sentence in sentences]

    p = document.add_paragraph(' ' * 7)

    for word_list in word_lists:
        for word in word_list:
            # 每一句的第一个单词大写
            if word_list.index(word) == 0 and word_lists.index(word_list) == 0:
                if word not in punctuation:
                    p.add_run(' ')
                    correct_word = correct_text_generic(word)
                    if correct_word != word:
                        correct_word = p.add_run(correct_word[0].upper() + correct_word[1:])
                        font = correct_word.font
                        font.color.rbg = RGBColor(0x00,0x00,0xff)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word)
                else:
                    p.add_run(word)
            else:
                p.add_run(' ')
                if word not in punctuation:
                    p.add_run(' ')
                    correct_word = correct_text_generic(word)
                    if correct_word != word:
                        correct_word = p.add_run(correct_word[0].upper() + correct_word[1:])
                        font = correct_word.font
                        font.color.rgb = RGBColor(0xff,0x00,0x00)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word)
                else:
                    p.add_run(word)

for i in range(len(file.paragraphs)):
    write_correct_paragraph(i)


document.save('../data/correct_document.1.docx')
print('修改完毕，一共修改了{}处'.format(COUNT_CORRECT))