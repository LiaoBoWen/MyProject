from docx import Document
from nltk import sent_tokenize,word_tokenize        # 对文章进行分割 ，对句子进行分割
from spell_correcter import correct_text_generic
from docx.shared import RGBColor
from string import punctuation
import nltk
# nltk.download('punkt')


# 文档中修改的个数
COUNT_CORRECT = 0

# 获取文档对象
file = Document('./data/Spelling Error.docx')

print('段落数量：',len(file.paragraphs))

punkt_list = r',.?"\'!()/\\-<>:@#$%^&*~'
# punkt_list = punctuation

document = Document()

def write_correct_paragraph(i):
    global COUNT_CORRECT
    # 每一段的内容
    paragraph = file.paragraphs[i].text.strip()
    # 对句子进行划分
    sentences = sent_tokenize(text=paragraph)
    # 词语划分
    words_list = [word_tokenize(sentence) for sentence in sentences]

    p = document.add_paragraph(' ' * 7)         # todo 这里是干什么
    for word_list in words_list:
        for word in word_list:
            # 对每一句话的第一个单词的第一个字母大写，并空两格
            if word_list.index(word) == 0 and words_list.index(word_list) == 0:
                if word not in punkt_list:
                    p.add_run(' ')
                    # 修改单词，如果单词正确，则返回单词
                    correct_word = correct_text_generic(word)
                    if correct_word != word:
                        colored_word = p.add_run(correct_word[0].upper() + correct_word[1:])
                        font = colored_word.font
                        font.color.rgb =  RGBColor(0x00,0x00,0xff)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word[0].upper() + correct_word[1:])
                else:
                    p.add_run(word)
            else:
                p.add_run(' ')
                # 修改单词，如果单词正确，这返回原单词
                correct_word = correct_text_generic(word)
                if word not in punkt_list :
                    # 如果单词改动，则单词变成红色
                    if correct_word != word:
                        colored_word = p.add_run(correct_word[0].upper() + correct_word[1:])
                        font = colored_word.font
                        font.color.rgb = RGBColor(0xff,0x00,0x00)
                        COUNT_CORRECT += 1
                    else:
                        p.add_run(correct_word)
                else:
                    p.add_run(word)

for i in range(len(file.paragraphs)):
    write_correct_paragraph(i)


document.save('./data/correct_document.docx')

print('修改并保存文件文完毕')
print('一共修改了{}处'.format(COUNT_CORRECT))
